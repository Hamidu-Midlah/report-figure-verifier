"""Deterministic tests for report ingestion (TXT, Markdown, DOCX, text PDF).

Fixtures are generated in-memory: DOCX via python-docx, and PDFs via a tiny
self-contained writer so no PDF-writing dependency (or committed binary) is
needed. Run with:  python -m tests.test_ingestion
"""

import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from verifier.ingestion import (
    extract_report_text,
    IngestionError,
    IngestionResult,
)
from verifier.tools import extract_numeric_claims


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------

def _docx_paragraphs(paragraphs):
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_with_table(rows, paragraphs=()):
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf(pages):
    """A minimal valid PDF. Each string is one page; "" means a no-text page."""
    def esc(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    objs = {}  # number -> bytes body
    objs[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objs[3] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    kids = []
    num = 4
    for text in pages:
        page_num, content_num = num, num + 1
        num += 2
        kids.append(f"{page_num} 0 R")
        stream = f"BT /F1 18 Tf 72 700 Td ({esc(text)}) Tj ET" if text else ""
        objs[content_num] = (
            f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream".encode("latin-1")
        )
        objs[page_num] = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {content_num} 0 R >>"
        ).encode("latin-1")
    objs[2] = (
        f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {len(pages)} >>"
    ).encode("latin-1")

    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for n in sorted(objs):
        offsets[n] = len(out)
        out += f"{n} 0 obj\n".encode("latin-1") + objs[n] + b"\nendobj\n"
    xref_pos = len(out)
    count = max(objs) + 1
    out += f"xref\n0 {count}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for n in range(1, count):
        out += f"{offsets[n]:010d} 00000 n \n".encode("latin-1")
    out += (
        f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF"
    ).encode("latin-1")
    return bytes(out)


def _figures(text):
    return {c["figure"] for c in extract_numeric_claims(text)}


EQUIV_TEXT = "In 2026, adoption reached 30% and satisfaction improved to 12%."
EQUIV_FIGURES = {"30%", "12%"}


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def test_txt_and_md_decode_and_preserve_paragraphs():
    body = "First paragraph with 40%.\n\nSecond paragraph with 12%."
    for ext in (".txt", ".md"):
        res = extract_report_text(body.encode("utf-8"), f"report{ext}")
        assert isinstance(res, IngestionResult)
        assert res.file_type == ext.lstrip(".")
        assert res.extracted_text == body            # unchanged text
        assert res.paragraph_count == 2
        assert _figures(res.extracted_text) == {"40%", "12%"}


def test_txt_invalid_utf8_is_rejected():
    try:
        extract_report_text(b"\xff\xfe bad bytes", "report.txt")
    except IngestionError as e:
        assert "UTF-8" in str(e)
    else:
        raise AssertionError("expected IngestionError for invalid UTF-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def test_docx_paragraphs():
    data = _docx_paragraphs(["In 2026, adoption reached 30%.",
                             "Satisfaction improved to 12%."])
    res = extract_report_text(data, "report.docx")
    assert res.file_type == "docx"
    assert res.paragraph_count == 2
    assert _figures(res.extracted_text) == {"30%", "12%"}


def test_docx_table_claims_distinguishable():
    data = _docx_with_table(
        rows=[["City", "2026 sales"], ["Manchester", "95k"], ["Leeds", "99k"]],
        paragraphs=["Regional sales table follows."],
    )
    res = extract_report_text(data, "report.docx")
    assert res.table_count == 1
    assert " | " in res.extracted_text  # cell separator preserved
    claims = extract_numeric_claims(res.extracted_text)
    figures = {c["figure"] for c in claims}
    assert "95k" in figures and "99k" in figures
    by_fig = {c["figure"]: c["sentence"] for c in claims}
    # each value stays with its own row, not merged across rows
    assert "Manchester" in by_fig["95k"] and "Leeds" in by_fig["99k"]
    assert "Leeds" not in by_fig["95k"]


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def test_pdf_text_extraction_and_page_markers():
    data = _make_pdf([EQUIV_TEXT])
    res = extract_report_text(data, "report.pdf")
    assert res.file_type == "pdf"
    assert res.page_count == 1
    assert "[Page 1]" in res.extracted_text        # boundary kept in preview
    assert _figures(res.extracted_text) == EQUIV_FIGURES  # page marker not a claim


def test_pdf_page_markers_never_become_claims():
    # A PDF whose only digits are the page markers must yield zero claims.
    data = _make_pdf(["No numbers appear on this page.",
                      "Only plain words here as well."])
    res = extract_report_text(data, "report.pdf")
    assert "[Page 1]" in res.extracted_text and "[Page 2]" in res.extracted_text
    assert extract_numeric_claims(res.extracted_text) == []


def test_pdf_isolated_blank_page_warns_but_keeps_others():
    data = _make_pdf(["Sales were 95k in early trading.", "",
                      "Revenue later rose to 99k."])
    res = extract_report_text(data, "report.pdf")
    assert res.page_count == 3
    assert res.extraction_status == "partial"
    assert any("pages 2" in w for w in res.warnings)
    figures = _figures(res.extracted_text)
    assert "95k" in figures and "99k" in figures   # surrounding pages survive


def test_pdf_fully_blank_is_reported_as_scanned():
    data = _make_pdf(["", ""])
    try:
        extract_report_text(data, "report.pdf")
    except IngestionError as e:
        assert "scanned or image-based" in str(e)
    else:
        raise AssertionError("expected scanned-PDF IngestionError")


def test_malformed_pdf_is_handled():
    try:
        extract_report_text(b"%PDF-1.4 this is not a real pdf body", "report.pdf")
    except IngestionError as e:
        assert "malformed" in str(e).lower() or "could not be read" in str(e).lower()
    else:
        raise AssertionError("expected IngestionError for malformed PDF")


# ---------------------------------------------------------------------------
# Cross-format equivalence and guards
# ---------------------------------------------------------------------------

def test_equivalent_content_across_formats():
    fixtures = {
        "report.txt": EQUIV_TEXT.encode("utf-8"),
        "report.md": EQUIV_TEXT.encode("utf-8"),
        "report.docx": _docx_paragraphs([EQUIV_TEXT]),
        "report.pdf": _make_pdf([EQUIV_TEXT]),
    }
    for name, data in fixtures.items():
        res = extract_report_text(data, name)
        assert _figures(res.extracted_text) == EQUIV_FIGURES, \
            f"{name}: {_figures(res.extracted_text)}"


def test_unsupported_file_type_is_rejected():
    try:
        extract_report_text(b"anything", "report.rtf")
    except IngestionError as e:
        assert "Unsupported" in str(e)
    else:
        raise AssertionError("expected IngestionError for unsupported type")


def test_preview_text_is_the_extraction_input():
    # The text shown in the preview is exactly the text fed to claim extraction:
    # one field, used verbatim by app.py for both preview and run_verification.
    data = _make_pdf(["Growth hit 30% in 2026.", "", "It fell to 24% later."])
    res = extract_report_text(data, "report.pdf")
    supplied = res.extracted_text
    assert "[Page 1]" in supplied and "[Page 2]" in supplied
    # claim extraction consumes exactly this string
    assert _figures(supplied) == {"30%", "24%"}


def main() -> int:
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_")]
    failed = 0
    for t in tests:
        try:
            t()
            print(f"PASS  {t.__name__}")
        except Exception as exc:  # noqa: BLE001
            failed += 1
            print(f"FAIL  {t.__name__}: {type(exc).__name__}: {exc}")
    print(f"\n{len(tests) - failed}/{len(tests)} tests passed")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
